from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import sys


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "chapters" / "00-table-of-contents.docx"
DOC_SKILL = Path(
    r"C:\Users\pmach\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
if DOC_SKILL.exists():
    sys.path.insert(0, str(DOC_SKILL))

from table_geometry import apply_table_geometry, column_widths_from_weights  # type: ignore

THAI_FONT = "TH Sarabun New"


def set_run_font(run, size: int | float | None = None, bold: bool | None = None):
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_para_font(paragraph, size: int = 16, bold: bool = False):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def setup_styles(doc: Document):
    styles = doc.styles
    for style_name, size, bold in [
        ("Normal", 16, False),
        ("Heading 1", 20, True),
        ("Heading 2", 18, True),
    ]:
        style = styles[style_name]
        style.font.name = THAI_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold


def add_toc_field(paragraph):
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = paragraph.add_run("คลิกขวาแล้วเลือก Update Field หลังรวมไฟล์รายงานทั้งหมด")
    set_run_font(placeholder, size=16)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.runs[-1]._r.append(end)


def add_heading(doc: Document, text: str):
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(p, size=20, bold=True)
    return p


def add_list_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["ลำดับ", "รายการ", "หน้า"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_font(p, size=16, bold=True)
    for no, title in rows:
        cells = table.add_row().cells
        values = [no, title, ""]
        for i, text in enumerate(values):
            cells[i].text = text
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_para_font(p, size=15)
    widths = column_widths_from_weights([1.1, 6.3, 1.0], 8765)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=0,
        cell_margins_dxa={"top": 95, "bottom": 95, "start": 130, "end": 130},
    )
    doc.add_paragraph()
    return table


def add_main_toc_table(doc: Document, rows: list[str]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["รายการ", "หน้า"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_font(p, size=16, bold=True)
    for title in rows:
        cells = table.add_row().cells
        cells[0].text = title
        cells[1].text = ""
        for i, cell in enumerate(cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_para_font(p, size=15)
    widths = column_widths_from_weights([7.2, 1.0], 8765)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=0,
        cell_margins_dxa={"top": 95, "bottom": 95, "start": 130, "end": 130},
    )
    doc.add_paragraph()
    return table


def build_doc(output_path: Path = OUTPUT):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)
    setup_styles(doc)

    add_heading(doc, "สารบัญ")
    main_rows = [
        "บทที่ 1 บทนำ",
        "1.1 ที่มาและความสำคัญของปัญหา",
        "1.2 ปัญหาของระบบงานเดิม",
        "1.3 วัตถุประสงค์ของโครงงาน",
        "1.4 ขอบเขตของโครงงาน",
        "1.5 ประโยชน์ที่คาดว่าจะได้รับ",
        "1.6 เครื่องมือและเทคโนโลยีที่ใช้",
        "1.7 ภาพรวมการทำงานของระบบ",
        "1.8 นิยามศัพท์เฉพาะ",
        "1.9 โครงสร้างรายงาน",
        "บทที่ 2 หลักการและทฤษฎีเบื้องต้น",
        "2.1 ข้อจำกัดของระบบค้นหาบุคคลแบบเดิมและแนวทางแก้ไข",
        "2.2 แนวคิดและทฤษฎีที่เกี่ยวข้องกับระบบ",
        "2.3 การเลือกใช้เทคโนโลยีในระบบ",
        "2.4 Dataset และการ Train โมเดล",
        "2.5 การจัดเก็บข้อมูลและระบบฐานข้อมูล",
        "2.6 การค้นหาและส่วนติดต่อผู้ใช้",
        "2.7 ข้อพิจารณาด้านกฎหมายและจริยธรรม",
        "บทที่ 3 วิธีดำเนินโครงงาน",
        "บทที่ 4 ผลการดำเนินงานและการทดสอบระบบ",
        "บทที่ 5 สรุปผลและข้อเสนอแนะ",
        "บรรณานุกรม",
        "ภาคผนวก",
    ]
    add_main_toc_table(doc, main_rows)
    doc.add_page_break()

    add_heading(doc, "สารบัญตาราง")
    table_rows = [
        ("ตารางที่ 1.1", "เครื่องมือและเทคโนโลยีหลักที่ใช้ในโครงงาน"),
        ("ตารางที่ 1.2", "ลำดับการทำงานหลักของระบบ"),
        ("ตารางที่ 2.1", "ประเภทเสื้อผ้า 6 คลาสที่ระบบตรวจจับ"),
        ("ตารางที่ 2.2", "โครงสร้างระบบสี 63 เฉดแบ่งตามกลุ่มสีหลัก"),
        ("ตารางที่ 2.3", "การเปรียบเทียบ DeepSORT และ ByteTrack"),
        ("ตารางที่ 2.4", "Dataset ที่ใช้ฝึกโมเดลตรวจจับเสื้อผ้า"),
        ("ตารางที่ 2.5", "Dataset จำนวนชิ้นของเสื้อผ้าที่นำมาฝึกและทดสอบโมเดล"),
        ("ตารางที่ 3.1", "คำอธิบายส่วนประกอบหลักของระบบ"),
        ("ตารางที่ 3.2", "ขอบเขตและข้อจำกัดของระบบ"),
        ("ตารางที่ 3.3", "ขั้นตอนการประมวลผลภาพ"),
        ("ตารางที่ 3.4", "เหตุผลในการเลือกโมเดลและ tracker"),
        ("ตารางที่ 3.5", "กลุ่ม class เสื้อผ้าที่ใช้ในระบบ"),
        ("ตารางที่ 3.6", "ประเภทข้อมูลสีในระบบ"),
        ("ตารางที่ 3.7", "ตารางหลักในฐานข้อมูล"),
        ("ตารางที่ 3.8", "โครงสร้างข้อมูลสำคัญของ detections"),
        ("ตารางที่ 3.9", "เครื่องมือและเวอร์ชันหลักจากโปรเจค"),
        ("ตารางที่ 3.10", "ตัวอย่าง API ที่ใช้ในระบบ"),
        ("ตารางที่ 3.11", "เครื่องมือการทดสอบ"),
        ("ตารางที่ 3.12", "แผนการทดสอบระบบ"),
        ("ตารางที่ 3.13", "ปัญหาและแนวทางแก้ไขระหว่างพัฒนา"),
        ("ตารางที่ 4.1", "ผลการเปรียบเทียบจำนวน detection ระหว่าง upload flow และ realtime flow"),
        ("ตารางที่ 4.2", "ผลการจำแนกประเภทเสื้อผ้าจากรายงาน parity test"),
        ("ตารางที่ 4.3", "ตัวอย่าง primary color group ที่ตรวจพบ"),
        ("ตารางที่ 4.4", "สภาพแวดล้อมในการทดสอบ"),
        ("ตารางที่ 4.5", "รูปแบบการทดสอบระบบ"),
        ("ตารางที่ 4.6", "ผลการทดสอบ flow parity"),
        ("ตารางที่ 4.7", "ปัญหาที่พบและแนวทางแก้ไข"),
        ("ตารางที่ 4.8", "ตัวอย่างเงื่อนไขการค้นหาที่ระบบรองรับ"),
        ("ตารางที่ 4.9", "ข้อมูลที่ควรเพิ่มเพื่อสรุปผลเชิงตัวเลข"),
    ]
    add_list_table(doc, table_rows)
    doc.add_page_break()

    add_heading(doc, "สารบัญภาพ")
    figure_rows = [
        ("ภาพที่ 3.1", "แผนภาพสถาปัตยกรรมส่วนประกอบของระบบ"),
        ("ภาพที่ 3.2", "แผนภาพกระบวนการประมวลผลแบบสองรอบ"),
        ("ภาพที่ 3.3", "แผนภาพการกู้คืน track id"),
        ("ภาพที่ 3.4", "หน้าจอ Dashboard"),
        ("ภาพที่ 3.5", "หน้าจอ Input Manager"),
        ("ภาพที่ 3.6", "หน้าจอ Realtime"),
        ("ภาพที่ 3.7", "หน้าจอ Search/Investigation"),
        ("ภาพที่ 3.8", "หน้าจอ Camera Management"),
        ("ภาพที่ 4.1", "หน้าจอ Dashboard ของระบบ"),
        ("ภาพที่ 4.2", "หน้าจอ Input Manager สำหรับอัปโหลดวิดีโอหรือเชื่อมต่อสตรีม"),
        ("ภาพที่ 4.3", "หน้าจอ Search/Investigation สำหรับค้นหาบุคคล"),
    ]
    add_list_table(doc, figure_rows)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = build_doc()
    print(out)
