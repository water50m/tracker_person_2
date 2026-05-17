from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "chapters" / "03-methodology-draft.md"
OUTPUT = ROOT / "chapters" / "03-methodology-chapter3.docx"
DOC_SKILL = Path(
    r"C:\Users\pmach\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
if DOC_SKILL.exists():
    sys.path.insert(0, str(DOC_SKILL))

from table_geometry import (  # type: ignore
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)

THAI_FONT = "TH Sarabun New"
LATIN_FONT = "TH Sarabun New"


def set_run_font(run, size: int | float | None = None, bold: bool | None = None):
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 14):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.strip())
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), "110")
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def set_table_width(table, width_dxa: int = 9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    set_run_font(run, size=12)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


def add_paragraph_with_inline_code(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(64, 64, 64)
        else:
            r = p.add_run(part)
            set_run_font(r, size=16)
    return p


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)


def parse_table(lines: list[str], start: int):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for raw in table_lines:
        cells = [clean_inline(c) for c in raw.strip("|").split("|")]
        if all(re.fullmatch(r"\s*:?-{3,}:?\s*", c) for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table)
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            set_cell_text(cell, text, bold=(ri == 0), size=13 if cols >= 4 else 14)
            if ri == 0:
                set_cell_shading(cell, "E8EEF7")
    content_width = section_content_width_dxa(doc.sections[0])
    if cols == 2:
        weights = [1.4, 3.6]
    elif cols == 3:
        weights = [1.35, 1.45, 3.2]
    elif cols == 4:
        weights = [1.25, 1.2, 1.7, 2.85]
    else:
        weights = [1] * cols
    widths = column_widths_from_weights(weights, content_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=0,
        cell_margins_dxa={"top": 95, "bottom": 95, "start": 130, "end": 130},
    )
    doc.add_paragraph()


def add_list_item(doc: Document, text: str, ordered: bool):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(clean_inline(text))
    set_run_font(r, size=16)


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = THAI_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
    normal.font.size = Pt(16)

    for name, size, color in [
        ("Title", 22, "000000"),
        ("Heading 1", 20, "1F4E79"),
        ("Heading 2", 18, "1F4E79"),
        ("Heading 3", 16, "333333"),
    ]:
        style = styles[name]
        style.font.name = THAI_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def build_doc(input_path: Path = INPUT, output_path: Path = OUTPUT):
    md = input_path.read_text(encoding="utf-8")
    lines = md.splitlines()
    header_label = next((line[2:].strip() for line in lines if line.startswith("# ")), "รายงานโครงงาน")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

    setup_styles(doc)

    header = section.header.paragraphs[0]
    header.text = ""
    hr = header.add_run(header_label)
    set_run_font(hr, size=12)
    hr.font.color.rgb = RGBColor(100, 100, 100)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(section.footer.paragraphs[0])

    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    i = 0
    title_seen = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip().strip("`")
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_buf))
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue

        if line.startswith("# "):
            text = clean_inline(line[2:])
            if title_seen:
                doc.add_page_break()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=22, bold=True)
            title_seen = True
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(clean_inline(line[3:]))
            set_run_font(r, size=20, bold=True)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(clean_inline(line[4:]))
            set_run_font(r, size=18, bold=True)
        elif re.match(r"^\d+\.\s+", line.strip()):
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", line.strip()), ordered=True)
        elif line.strip().startswith("- "):
            add_list_item(doc, line.strip()[2:], ordered=False)
        elif line.strip() == "---":
            p = doc.add_paragraph()
            r = p.add_run("─" * 45)
            set_run_font(r, size=12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Convert draft guidance phrases into final document phrasing.
            text = line.strip()
            text = text.replace("ภาพที่ 3.1 ควรแสดงสถาปัตยกรรมโดยรวมของระบบ โดยประกอบด้วย", "ภาพที่ 3.1 แสดงสถาปัตยกรรมโดยรวมของระบบ ซึ่งประกอบด้วย")
            text = text.replace("ภาพที่ 3.2 ควรแสดงลำดับการทำงานตั้งแต่รับเฟรมจนบันทึกผลลัพธ์", "ภาพที่ 3.2 แสดงลำดับการทำงานตั้งแต่รับเฟรมจนบันทึกผลลัพธ์")
            text = text.replace("ภาพที่ 3.3 ควรแสดง flow การกู้คืน track id", "ภาพที่ 3.3 แสดงขั้นตอนการกู้คืน track id")
            add_paragraph_with_inline_code(doc, text)
        i += 1

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default=str(INPUT))
    parser.add_argument("--output", default=str(OUTPUT))
    args = parser.parse_args()
    out = build_doc(Path(args.input), Path(args.output))
    print(out)
